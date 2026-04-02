# ============================================================================
# GENERADOR DE INFORME - TAREA 1 DISEÑO SISMORRESISTENTE
# Genera el documento Word con carátula PUCE, figuras e interpretaciones
# ============================================================================

import os
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FIGURAS_DIR = os.path.join(BASE_DIR, 'figuras')
OUTPUT_PATH = os.path.join(BASE_DIR, 'Tarea1_DisenoSismorresistente_PabloBaez.docx')


def crear_documento():
    """Crea el documento Word con formato PUCE."""
    doc = Document()

    # Configurar página A4 con márgenes de 2.5 cm
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Configurar estilo Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 2.0
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    return doc


def agregar_caratula(doc):
    """Agrega la carátula institucional PUCE."""
    # Líneas vacías superiores
    for _ in range(2):
        doc.add_paragraph('')

    # Universidad
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PONTIFICIA UNIVERSIDAD CATÓLICA DEL ECUADOR')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph('')

    # Escuela
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ESCUELA DE ARQUITECTURA, DISEÑO Y ARTES\nCARRERA DE INGENIERÍA CIVIL')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    for _ in range(2):
        doc.add_paragraph('')

    # Periodo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PERIODO ACADÉMICO:')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PRIMER PERIODO 2026')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph('')

    # Tema
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TEMA:')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ANÁLISIS DEL CATÁLOGO SÍSMICO HOMOGENEIZADO DEL ECUADOR')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph('')

    # Materia
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('MATERIA:')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DISEÑO SISMORRESISTENTE')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph('')

    # Nivel
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('NIVEL/PARALELO:')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('8VO. "A"')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph('')

    # Alumno
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ALUMNO:')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PABLO BAEZ')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph('')

    # Docente
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DOCENTE:')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('MSc. MIGUEL RIVERA')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph('')

    # Fecha
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('FECHA DE ENTREGA:')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ABRIL 2026')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    # Salto de página
    doc.add_page_break()


def agregar_titulo(doc, texto):
    """Agrega un título nivel 1 (centrado, negrita)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'


def agregar_subtitulo(doc, texto):
    """Agrega un título nivel 2 (izquierda, negrita)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'


def agregar_parrafo(doc, texto):
    """Agrega un párrafo justificado con sangría francesa."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(texto)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'


def agregar_imagen(doc, nombre_archivo, ancho_cm=14):
    """Agrega una imagen centrada."""
    ruta = os.path.join(FIGURAS_DIR, nombre_archivo)
    if os.path.exists(ruta):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(ruta, width=Cm(ancho_cm))
    else:
        agregar_parrafo(doc, f'[Imagen no encontrada: {nombre_archivo}]')


def agregar_pie_figura(doc, texto):
    """Agrega pie de figura centrado en cursiva."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(texto)
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'


def construir_informe():
    """Construye el informe completo."""
    print("Generando informe Word...")
    doc = crear_documento()

    # ==================== CARÁTULA ====================
    agregar_caratula(doc)

    # ==================== INTRODUCCIÓN ====================
    agregar_titulo(doc, 'Introducción')

    agregar_parrafo(doc,
        'El presente informe desarrolla el análisis del catálogo sísmico '
        'homogeneizado del Ecuador, elaborado por el Instituto Geofísico de la '
        'Escuela Politécnica Nacional (IGEPN) y documentado en la publicación de '
        'Beauval et al. (2013). El catálogo abarca el período comprendido entre '
        '1901 y 2009, conteniendo 10.742 eventos sísmicos con magnitud Mw igual '
        'o superior a 3.0, cuyas magnitudes fueron homogeneizadas a la escala de '
        'magnitud momento (Mw). A partir de estos datos se realizan análisis '
        'espaciales, estadísticos y de frecuencia que permiten caracterizar el '
        'comportamiento de la sismicidad en el territorio ecuatoriano y su '
        'relación con el alto peligro sísmico que enfrenta el país, debido '
        'principalmente a la subducción de la Placa de Nazca bajo la Placa '
        'Sudamericana y la presencia de fallas activas en la región interandina.'
    )

    # ==================== LITERAL 1 ====================
    agregar_titulo(doc, '1. Análisis Espacial de Eventos Sísmicos')

    agregar_parrafo(doc,
        'Para el análisis espacial se elaboró un mapa de dispersión de los '
        'eventos sísmicos registrados en el catálogo, utilizando la magnitud Mw '
        'como variable de referencia. Se diferenciaron tres rangos de magnitud: '
        'eventos con Mw mayor a 7 (representados en rojo con marcadores grandes), '
        'eventos con Mw entre 5 y 7 (en naranja con marcadores medianos) y '
        'eventos con Mw entre 4 y 5 (en amarillo con marcadores pequeños). Los '
        'eventos con magnitud inferior a 3 fueron excluidos del análisis, tal '
        'como lo establece el enunciado de la tarea.'
    )

    agregar_imagen(doc, 'literal1_mapa_sismico.png', ancho_cm=15)
    agregar_pie_figura(doc, 'Figura 1. Mapa de eventos sísmicos en el Ecuador, '
                       'Catálogo Homogeneizado (1901-2009).')

    agregar_parrafo(doc,
        'El mapa muestra una concentración notable de eventos sísmicos a lo '
        'largo de la zona costera, particularmente en la región norte '
        '(Esmeraldas y Manabí), lo cual se asocia directamente con la zona de '
        'subducción de la Placa de Nazca. Los eventos de mayor magnitud (Mw > 7) '
        'se distribuyen principalmente en esta franja costera, confirmando que la '
        'interfaz de subducción constituye la fuente sismogénica más peligrosa '
        'del país. En la región interandina (Sierra) se observa una sismicidad '
        'moderada pero persistente, asociada al sistema de fallas activas que '
        'atraviesa el callejón interandino. En total se registraron 2.581 eventos '
        'con Mw entre 4 y 5, 534 eventos con Mw entre 5 y 7, y 22 eventos con '
        'Mw superior a 7.'
    )

    agregar_parrafo(doc,
        'Adicionalmente se generó un mapa de calor interactivo mediante la '
        'librería Folium, el cual permite una exploración dinámica de la '
        'distribución espacial de los eventos sísmicos. Este mapa se entrega '
        'como archivo HTML adjunto.'
    )

    # ==================== LITERAL 2 ====================
    agregar_titulo(doc, '2. Identificación de Eventos Relevantes')

    agregar_subtitulo(doc, 'Evento sísmico más antiguo')

    agregar_parrafo(doc,
        'El evento sísmico más antiguo registrado en el catálogo homogeneizado '
        'data del 7 de enero de 1901, con una magnitud Mw de 7.2, ubicado en '
        'las coordenadas latitud -2.00° y longitud -82.00°, correspondiente a '
        'la zona oceánica frente a la costa ecuatoriana. Este registro proviene '
        'del catálogo CENTENNIAL con fuente ABE, lo que evidencia que el Ecuador '
        'cuenta con registros sísmicos instrumentales desde inicios del siglo XX, '
        'complementados con información de catálogos internacionales.'
    )

    agregar_subtitulo(doc, 'Evento sísmico de mayor magnitud')

    agregar_parrafo(doc,
        'El evento de mayor magnitud registrado en el catálogo corresponde al '
        'terremoto de Esmeraldas del 31 de enero de 1906, con una magnitud Mw '
        'de 8.35, ubicado en las coordenadas latitud 0.95° y longitud -79.36°, '
        'a una profundidad de 20.0 km. Este sismo ocurrió en la provincia de '
        'Esmeraldas, en la zona de subducción de la costa norte del Ecuador, y '
        'es considerado uno de los terremotos más grandes registrados a nivel '
        'mundial en el siglo XX. El evento generó un tsunami que afectó la costa '
        'del Pacífico suramericano.'
    )

    agregar_subtitulo(doc, 'Eventos de mayor magnitud registrados')

    agregar_parrafo(doc,
        'La siguiente tabla presenta los 10 eventos sísmicos de mayor magnitud '
        'registrados en el catálogo homogeneizado del Ecuador:'
    )

    # Tabla de top 10 eventos
    tabla = doc.add_table(rows=11, cols=7)
    tabla.style = 'Light Grid Accent 1'
    encabezados = ['N°', 'Año', 'Fecha', 'Latitud', 'Longitud', 'Prof. (km)', 'Mw']
    for i, enc in enumerate(encabezados):
        cell = tabla.rows[0].cells[i]
        cell.text = enc
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'

    datos_top10 = [
        ['1', '1906', '31/01', '0.95', '-79.36', '20.0', '8.35'],
        ['2', '1979', '12/12', '1.62', '-79.42', '37.5', '8.10'],
        ['3', '1942', '14/05', '0.01', '-80.39', '19.7', '7.80'],
        ['4', '1958', '19/01', '1.14', '-79.59', '29.8', '7.80'],
        ['5', '1960', '20/11', '-6.77', '-80.91', '15.0', '7.80'],
        ['6', '1901', '07/01', '-2.00', '-82.00', '0.0', '7.20'],
        ['7', '1928', '14/05', '-5.26', '-78.55', '35.0', '7.20'],
        ['8', '1906', '28/09', '-2.00', '-79.00', '150.0', '7.50'],
        ['9', '1917', '31/08', '4.00', '-74.00', '0.0', '7.10'],
        ['10', '1907', '01/06', '0.00', '-82.00', '0.0', '7.00'],
    ]
    for fila_idx, fila_datos in enumerate(datos_top10):
        for col_idx, valor in enumerate(fila_datos):
            cell = tabla.rows[fila_idx + 1].cells[col_idx]
            cell.text = valor
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'

    agregar_pie_figura(doc, 'Tabla 1. Eventos sísmicos de mayor magnitud del catálogo '
                       'homogeneizado del Ecuador (1901-2009).')

    agregar_parrafo(doc,
        'Los 10 eventos de mayor magnitud se concentran en la zona costera norte '
        'del Ecuador, particularmente en la región de Esmeraldas y Manabí, '
        'confirmando que la interfaz de subducción entre la Placa de Nazca y la '
        'Placa Sudamericana es la principal fuente de peligro sísmico del país. '
        'Varios de estos eventos ocurrieron en coordenadas muy similares '
        '(latitud aproximada 1°N, longitud entre -79° y -80°W), lo que sugiere '
        'una recurrencia de grandes terremotos en este segmento de la zona de '
        'subducción, con períodos de retorno que deben considerarse en el diseño '
        'sismorresistente de estructuras en la región.'
    )

    # ==================== LITERAL 3 ====================
    agregar_titulo(doc, '3. Análisis de Frecuencias por Intervalos de Magnitud')

    agregar_parrafo(doc,
        'Se elaboró un histograma de frecuencias para los eventos sísmicos con '
        'Mw igual o superior a 4.0, clasificados en siete intervalos de magnitud '
        'definidos por el enunciado de la tarea: 4.0-4.5, 4.5-5.0, 5.0-5.5, '
        '5.5-6.0, 6.0-6.5, 6.5-7.0 y mayores a 7.0.'
    )

    agregar_imagen(doc, 'literal3_histograma_magnitudes.png', ancho_cm=14)
    agregar_pie_figura(doc, 'Figura 2. Histograma de frecuencias de eventos sísmicos '
                       'por intervalo de magnitud.')

    agregar_parrafo(doc,
        'El histograma muestra una distribución inversamente proporcional entre '
        'la magnitud y la frecuencia de eventos, lo cual es consistente con la '
        'ley de Gutenberg-Richter (log N = a - bM). El intervalo de 4.0 a 4.5 '
        'concentra la mayor cantidad de eventos con 1.499 registros, lo que '
        'representa el 47.7% del total de sismos con Mw igual o superior a 4.0. '
        'A medida que aumenta la magnitud, la frecuencia disminuye de forma '
        'exponencial: 1.082 eventos en el rango 4.5-5.0, 325 en el rango '
        '5.0-5.5, 103 en el rango 5.5-6.0, 60 en el rango 6.0-6.5, 46 en el '
        'rango 6.5-7.0, y apenas 22 eventos con magnitud superior a 7.0. Esta '
        'distribución refleja el comportamiento típico de una zona de alta '
        'sismicidad asociada a un margen de subducción activo, donde los eventos '
        'de gran magnitud son poco frecuentes pero representan el mayor peligro '
        'para las estructuras y la población.'
    )

    # ==================== LITERAL 4 ====================
    agregar_titulo(doc, '4. Análisis Estadístico de Magnitudes')

    agregar_parrafo(doc,
        'Se calcularon las estadísticas descriptivas de las magnitudes de los '
        '10.742 eventos sísmicos del catálogo. La magnitud promedio es de 3.69 '
        'Mw, la mediana es de 3.40 Mw y la desviación estándar es de 0.70. Se '
        'ajustó una distribución normal a los datos de magnitud mediante el '
        'método de máxima verosimilitud, obteniéndose los cuantiles P16 = 2.99 '
        'y P84 = 4.38, los cuales se representan mediante líneas verticales en '
        'la gráfica junto con el promedio y la mediana.'
    )

    agregar_imagen(doc, 'literal4_distribucion_magnitudes.png', ancho_cm=14)
    agregar_pie_figura(doc, 'Figura 3. Distribución normal ajustada a las magnitudes '
                       'sísmicas con indicación del promedio, P16 y P84.')

    agregar_parrafo(doc,
        'La mediana (3.40) es menor que la media (3.69), lo que indica una '
        'distribución con sesgo positivo: la presencia de eventos de gran '
        'magnitud desplaza el promedio hacia valores más altos. El intervalo '
        'entre P16 (2.99) y P84 (4.38) contiene aproximadamente el 68% de los '
        'eventos, lo que equivale a una desviación estándar en torno a la media '
        'en una distribución normal. La distribución de magnitudes no se ajusta '
        'perfectamente a una normal, ya que los catálogos sísmicos siguen una '
        'distribución exponencial truncada según la ley de Gutenberg-Richter. '
        'No obstante, el ajuste normal proporciona una aproximación estadística '
        'útil para describir la tendencia central y la dispersión de los datos.'
    )

    # ==================== LITERAL 5 ====================
    agregar_titulo(doc, '5. Análisis Estadístico de Profundidades')

    agregar_parrafo(doc,
        'Se analizaron las profundidades focales de los eventos sísmicos del '
        'catálogo. La profundidad promedio es de 40.2 km, la mediana es de 18.2 '
        'km y la desviación estándar es de 49.4 km. Se ajustó una distribución '
        'normal a los datos de profundidad, obteniéndose los cuantiles P16 = 0.0 '
        'km (ajustado, ya que el valor teórico resulta negativo) y P84 = 89.3 km.'
    )

    agregar_imagen(doc, 'literal5_distribucion_profundidades.png', ancho_cm=14)
    agregar_pie_figura(doc, 'Figura 4. Distribución normal ajustada a las profundidades '
                       'sísmicas con indicación del promedio, P16 y P84.')

    agregar_parrafo(doc,
        'La mediana (18.2 km) es considerablemente menor que la media (40.2 km), '
        'lo que indica una distribución fuertemente sesgada hacia profundidades '
        'someras: la mayoría de los sismos ocurren a profundidades relativamente '
        'superficiales. Esto es coherente con el contexto tectónico del Ecuador, '
        'donde los sismos corticales (0-30 km), asociados a las fallas activas '
        'de la Sierra, y los sismos de la interfaz de subducción (30-60 km) son '
        'los más frecuentes. Los sismos profundos (mayores a 200 km) corresponden '
        'a la zona de Wadati-Benioff, donde la Placa de Nazca se hunde bajo el '
        'continente sudamericano. La distribución normal no se ajusta de manera '
        'ideal a los datos de profundidad, dado que las profundidades focales '
        'presentan una distribución asimétrica por naturaleza, sin embargo '
        'permite identificar las tendencias centrales y la dispersión de los datos.'
    )

    # ==================== LITERAL 6 ====================
    agregar_titulo(doc, '6. Conclusiones')

    agregar_parrafo(doc,
        'El análisis del catálogo sísmico homogeneizado del Ecuador (1901-2009) '
        'permite establecer las siguientes conclusiones técnicas respecto al '
        'comportamiento de la sismicidad y su relación con el peligro sísmico '
        'del territorio ecuatoriano:'
    )

    agregar_parrafo(doc,
        'El Ecuador presenta una sismicidad elevada y sostenida, con 10.742 '
        'eventos sísmicos registrados con Mw igual o superior a 3.0 en un '
        'período de 109 años, incluyendo 22 eventos de magnitud Mw igual o '
        'superior a 7.0. Esta actividad sísmica se asocia directamente con la '
        'subducción de la Placa de Nazca bajo la Placa Sudamericana, proceso '
        'tectónico que genera los sismos de mayor magnitud en la zona costera '
        'del país, y con la presencia de fallas activas en la región interandina.'
    )

    agregar_parrafo(doc,
        'Los eventos de mayor magnitud se concentran en la zona costera norte, '
        'particularmente en la región de Esmeraldas, donde el terremoto de 1906 '
        '(Mw 8.35) constituye el evento más grande registrado. La recurrencia '
        'de grandes terremotos en este segmento de subducción, evidenciada por '
        'los eventos de 1906, 1942, 1958 y 1979, confirma la necesidad de '
        'considerar escenarios de amenaza sísmica severa en el diseño '
        'sismorresistente de las estructuras de la región.'
    )

    agregar_parrafo(doc,
        'La predominancia de sismos superficiales (profundidad mediana de 18.2 '
        'km) implica que los efectos en superficie son más severos, ya que la '
        'atenuación de las ondas sísmicas es menor a profundidades someras. La '
        'distribución de frecuencias por magnitud sigue el patrón esperado de '
        'Gutenberg-Richter, lo que valida la consistencia del catálogo para su '
        'uso en análisis de peligro sísmico probabilístico.'
    )

    agregar_parrafo(doc,
        'Estos resultados justifican la clasificación del Ecuador como un país '
        'de alto peligro sísmico, tal como lo establece la Norma Ecuatoriana de '
        'la Construcción NEC-SE-DS, y refuerzan la importancia de aplicar '
        'criterios rigurosos de diseño sismorresistente en todas las '
        'edificaciones del territorio nacional.'
    )

    # ==================== REFERENCIAS ====================
    doc.add_page_break()
    agregar_titulo(doc, 'Referencias')

    referencias = [
        'Beauval, C., Yepes, H., Palacios, P., Segovia, M., Alvarado, A., '
        'Font, Y., Aguilar, J., Troncoso, L. y Vaca, S. (2013). An Earthquake '
        'Catalog for Seismic Hazard Assessment in Ecuador. Bulletin of the '
        'Seismological Society of America, 103(2A), 773-786. '
        'https://doi.org/10.1785/0120120270',

        'Instituto Geofísico de la Escuela Politécnica Nacional [IGEPN]. (2026). '
        'Catálogo sísmico homogeneizado del Ecuador. Descargado el 1 de abril '
        'de 2026 de https://www.igepn.edu.ec/catalogos-sismicos',

        'Norma Ecuatoriana de la Construcción [NEC]. (2015). NEC-SE-DS: Peligro '
        'Sísmico y Diseño Sismo Resistente. Ministerio de Desarrollo Urbano '
        'y Vivienda del Ecuador.',
    ]

    for ref in referencias:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 2.0
        # Sangría francesa: primera línea 0, resto con sangría izquierda
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(ref)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    # Guardar
    doc.save(OUTPUT_PATH)
    print(f"Informe guardado en: {OUTPUT_PATH}")
    print("Listo. Revisa el documento y expórtalo a PDF desde Word.")


if __name__ == '__main__':
    construir_informe()
